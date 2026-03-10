"""DOCX file parser (requires python-docx)."""

from __future__ import annotations

from ragmax.core.models import Document
from ragmax.core.utils import require_dependency


class DocxParser:
    """Extract text from Word .docx files."""

    def __init__(self) -> None:
        require_dependency("docx", "parsers")

    def parse(self, path: str) -> Document:
        import docx

        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return Document(
            content="\n\n".join(paragraphs),
            source=path,
            metadata={"parser": "docx"},
        )

    def supports(self, path: str) -> bool:
        return path.lower().endswith(".docx")
